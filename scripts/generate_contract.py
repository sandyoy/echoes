#!/usr/bin/env python3
"""生成数字疗法推广服务协议 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ========== 样式设置 ==========
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)

def add_run(paragraph, text, bold=False, size=12):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return run

# ========== 标题 ==========
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(title, '数字疗法推广服务协议', bold=True, size=22)

doc.add_paragraph()

# ========== 甲乙双方 ==========
def add_party(label, company, items):
    p = doc.add_paragraph()
    add_run(p, f'{label}：{company}', bold=True)
    for key, val in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        add_run(p, f'{key}：{val}')

add_party('甲  方', '湖南万物成理医疗科技有限公司', [
    ('统一社会信用代码', '（待补充）'),
    ('法定代表人', '（待补充）'),
    ('注册地址', '（待补充）'),
])

doc.add_paragraph()

add_party('乙  方', '圆爱康生物科技有限公司', [
    ('统一社会信用代码', '91440101MA5CBEEG5D'),
    ('法定代表人', '欧阳日红'),
    ('注册地址', '广州市天河区高普路38号7层713室'),
])

doc.add_paragraph()

# ========== 序言 ==========
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0.75)
add_run(p, '鉴于甲方拥有数字疗法产品的合法经营资质及市场推广需求，'
        '乙方拥有基层医疗机构渠道资源和社区健康服务推广能力。'
        '甲乙双方本着平等互利、诚实信用的原则，就数字疗法产品在'
        '圆爱康·智慧家康平台的推广服务事宜，经友好协商，达成如下协议：')

# ========== 条款 ==========
clauses = [
    ('第一条  定义与服务内容',
     '1.1 数字疗法产品：指甲方拥有合法知识产权或经营授权，通过软件程序驱动的、'
     '基于循证医学（Evidence-based Medicine）的治疗干预方案，'
     '包括但不限于注意缺陷多动障碍（ADHD）数字疗法、盆底康复数字疗法、'
     '心脏术后康复数字疗法等产品。\n\n'
     '1.2 推广服务：乙方通过圆爱康·智慧家康平台，与各级医疗机构合作，'
     '向有康复需求的终端患者推广数字疗法产品。'
     '推广方式包括但不限于医疗机构推荐、平台引导、社区健康服务站点体验转化等。\n\n'
     '1.3 推广区域：中华人民共和国境内（不含港澳台地区）。'),

    ('第二条  合作期限',
     '2.1 本协议有效期为三年，自双方签署之日起生效。\n\n'
     '2.2 协议期满前六十日内，如双方均未提出书面终止通知，本协议自动续签一年，'
     '以此类推。\n\n'
     '2.3 任何一方可在提前三十日书面通知对方的情况下解除本协议。'),

    ('第三条  计费方式及结算',
     '【收益参考说明】\n\n'
     '为便于理解本协议约定的技术服务费所产生的收益规模，举例如下：\n'
     '一个患者完成一套完整的康复训练课程，'
     '甲方按约定可收取约75元的技术服务费。以此推算，在乙方推广的支持下，'
     '甲方通过圆爱康·智慧家康平台可获得的月度收益参考如下：\n'
     '  - 月推广 500 例 → 甲方收益约 37,500 元\n'
     '  - 月推广 1,000 例 → 甲方收益约 75,000 元\n'
     '  - 月推广 2,000 例 → 甲方收益约 150,000 元\n'
     '上述收益为甲方直接收入，乙方向甲方支付的技术服务费5元/次'
     '在该收益结构中占比极小，双方合作具有显著的共赢空间。\n\n'
     '3.1 费用标准：乙方按推广实际例数向甲方支付技术服务费，'
     '标准为人民币5元/次。\n\n'
     '3.2 "一次"的定义：一个患者通过乙方推广渠道在合作医院'
     '激活并使用指定数字疗法产品，以第一个完整使用日为计费节点。\n\n'
     '3.3 结算周期：月度结算。当月消耗数据，次月15日前'
     '双方完成对账确认。乙方在收到甲方开具的合法有效发票后'
     '10个工作日内完成付款。\n\n'
     '3.4 对账方式：甲乙双方以双方确认的后台系统数据为准进行对账。'
     '如有争议，以双方共同确认的数据为准。\n\n'
     '3.5 本协议约定的技术服务费为甲乙双方之间独立的费用结算，'
     '不涉及任何第三方的交易结构或定价安排。任何一方不得以与第三方之间的合作条件为由'
     '要求变更本协议约定的费用标准。\n'),

    ('第四条  双方权利与义务',
     '4.1 甲方权利与义务：\n'
     '  （1）甲方应保证其数字疗法产品符合国家相关法律法规要求，'
     '取得必要的注册证、备案凭证或资质许可；\n'
     '  （2）甲方应为乙方提供必要的产品培训材料、宣传资料和技术支持；\n'
     '  （3）甲方应按约定及时足额支付推广服务费；\n'
     '  （4）甲方有权对乙方的推广行为进行监督，发现违规推广行为有权要求乙方立即纠正。\n\n'
     '4.2 乙方权利与义务：\n'
     '  （1）乙方应按照甲方提供的产品信息和推广规范进行推广，不得夸大宣传、虚假宣传；\n'
     '  （2）乙方应维护甲方及产品的品牌形象和商业信誉；\n'
     '  （3）乙方有义务向合作医疗机构说明产品的使用方法和注意事项；\n'
     '  （4）乙方应按约定向甲方提供推广数据和患者使用数据。'),

    ('第五条  数据权益',
     '5.1 原始生理数据归属：通过本协议约定的数字疗法设备采集的原始生理数据'
     '（包括但不限于心率、肌电信号、盆底肌收缩力等基础生物信号数据），'
     '归设备提供方（乙方）所有。乙方有权在遵守个人信息保护法律法规的前提下，'
     '将上述数据用于产品优化、算法迭代与学术研究。\n\n'
     '5.2 诊疗数据归属：基于原始数据生成的诊疗分析报告、诊断结论及治疗方案，'
     '归医疗机构（合作医院）所有。\n\n'
     '5.3 患者个人信息归属：患者个人信息归患者本人所有。甲乙双方及合作医院'
     '使用患者个人信息时，须获得患者书面授权，并严格遵守'
     '《中华人民共和国个人信息保护法》《中华人民共和国数据安全法》及相关法律法规的规定。\n\n'
     '5.4 数据使用限制：任何一方将数据提供给第三方时，'
     '须获得该数据权利人的书面同意。'),

    ('第六条  保密义务',
     '6.1 双方对在合作过程中知悉的对方商业秘密'
     '（包括但不限于产品信息、客户信息、定价策略、技术资料等）负有保密义务。'
     '保密义务不因本协议的终止而解除。\n\n'
     '6.2 保密信息不包括：\n'
     '  （1）在披露时已为公众所知的信息；\n'
     '  （2）接收方能够证明在披露前已知悉的信息；\n'
     '  （3）从有权披露的第三方合法获取的信息。'),

    ('第七条  违约责任',
     '7.1 任何一方违反本协议约定，给对方造成损失的，应承担相应的违约赔偿责任。\n\n'
     '7.2 甲方逾期支付服务费的，每逾期一日，应按逾期金额的万分之五向乙方支付违约金。\n\n'
     '7.3 因乙方违规推广导致甲方被行政处罚或遭受第三方索赔的，'
     '乙方应承担相应的赔偿责任。'),

    ('第八条  不可抗力',
     '8.1 因自然灾害、战争、政府行为、政策调整、重大公共卫生事件等不可抗力因素'
     '导致本协议无法履行的，受影响方不承担违约责任，但应及时通知对方并提供相关证明。\n\n'
     '8.2 不可抗力持续六十日以上的，任何一方有权书面通知对方终止本协议。'),

    ('第九条  争议解决',
     '9.1 本协议的订立、解释、履行及争议解决均适用中华人民共和国法律。\n\n'
     '9.2 因本协议引起的或与本协议有关的任何争议，双方应首先通过友好协商解决；'
     '协商不成的，任何一方均有权向乙方所在地有管辖权的人民法院提起诉讼。'),

    ('第十条  其他',
     '10.1 本协议未尽事宜，双方可另行签订补充协议，补充协议与本协议具有同等法律效力。\n\n'
     '10.2 本协议一式两份，甲乙双方各执一份，具有同等法律效力。\n\n'
     '10.3 本协议自双方签字（或盖章）之日起生效。\n\n'
     '10.4 双方确认，本合同记载的地址为有效送达地址。'
     '如任何一方地址变更，应自变更之日起三日内书面通知对方。'),
]

for title_text, body_text in clauses:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    add_run(p, title_text, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.75)
    add_run(p, body_text)

# ========== 签署页 ==========
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, '（以下为签章页，无正文）')

doc.add_paragraph()
doc.add_paragraph()

table = doc.add_table(rows=2, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 甲方列
items_a = [
    ('甲方：湖南万物成理医疗科技有限公司', True),
    ('法定代表人（或授权代表）：', False),
    ('日期：      年    月    日', False),
    ('联系电话：', False),
    ('地址：', False),
]
for i, (text, bold) in enumerate(items_a):
    if i == 0:
        p = table.cell(0, 0).paragraphs[0]
    else:
        p = table.cell(0, 0).add_paragraph()
    add_run(p, text, bold=bold)

# 乙方列
items_b = [
    ('乙方：圆爱康生物科技有限公司', True),
    ('法定代表人（或授权代表）：', False),
    ('日期：      年    月    日', False),
    ('联系电话：', False),
    ('地址：广州市天河区高普路38号7层713室', False),
]
for i, (text, bold) in enumerate(items_b):
    if i == 0:
        p = table.cell(0, 1).paragraphs[0]
    else:
        p = table.cell(0, 1).add_paragraph()
    add_run(p, text, bold=bold)

for row in table.rows:
    row.cells[0].width = Cm(9)
    row.cells[1].width = Cm(9)

# ========== 保存 ==========
output_dir = '/home/ubuntu/echoes/docs'
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, '数字疗法推广服务协议.docx')
doc.save(output_path)
print(f'✅ 合同已生成：{output_path}')
print(f'📏 文件大小：{os.path.getsize(output_path)/1024:.1f} KB')
