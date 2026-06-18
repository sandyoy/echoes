#!/usr/bin/env python3
"""Convert Markdown to properly formatted Word document using python-docx."""

import re
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def add_bold_to_paragraph(paragraph, text):
    """Add text to paragraph, processing **bold** markers."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)

def set_heading_style(doc, level, text):
    """Add a heading with consistent style."""
    if level == 1:
        heading = doc.add_heading(text, level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
            run.font.size = Pt(22)
    elif level == 2:
        heading = doc.add_heading(text, level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
            run.font.size = Pt(16)
    elif level == 3:
        heading = doc.add_heading(text, level=3)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x2D, 0x5F, 0x9E)
            run.font.size = Pt(13)
    else:
        heading = doc.add_heading(text, level=4)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            run.font.size = Pt(12)
    return heading

def add_normal_paragraph(doc, text):
    """Add a normal paragraph with bold support."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    add_bold_to_paragraph(p, text)
    return p

def add_bullet_point(doc, text, level=0):
    """Add a bullet point with bold support."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.8)
    add_bold_to_paragraph(p, text)
    return p

def add_code_block(doc, code_text):
    """Add a code block with monospace font and gray background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0"/>')
    p.paragraph_format.element.get_or_add_pPr().append(shading)
    return p

def add_blockquote(doc, text):
    """Add a blockquote (indented, left border, gray italic)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    pPr = p.paragraph_format.element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:left w:val="single" w:sz="12" w:space="8" w:color="1A478A"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    return p

def add_table_from_md(doc, lines, start_idx):
    """Parse and add a markdown table."""
    header_line = lines[start_idx]
    headers = [h.strip() for h in header_line.split('|')[1:-1]]
    
    rows_data = []
    i = start_idx + 2
    while i < len(lines):
        line = lines[i].strip()
        if not line or not line.startswith('|'):
            break
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows_data.append(cells)
        i += 1
    
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A478A" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    
    for r_idx, row in enumerate(rows_data):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            add_bold_to_paragraph(p, cell_text)
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F7FA" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    
    doc.add_paragraph()
    return i

def convert_md_to_docx(md_path, docx_path):
    """Main conversion function."""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    for level in range(1, 5):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    i = 0
    in_code_block = False
    code_buffer = []
    
    while i < len(lines):
        line = lines[i]
        
        if line.strip().startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                i += 1
                continue
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue
        
        if not line.strip():
            i += 1
            continue
        
        h_match = re.match(r'^(#{1,4})\s+(.+)$', line)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2).strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            set_heading_style(doc, level, text)
            i += 1
            continue
        
        if re.match(r'^---+$', line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('─' * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            i += 1
            continue
        
        if line.strip().startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            i = add_table_from_md(doc, lines, i)
            continue
        
        if line.strip().startswith('>'):
            text = re.sub(r'^>\s*', '', line)
            add_blockquote(doc, text)
            i += 1
            continue
        
        if re.match(r'^[\s]*[-*]\s+', line):
            text = re.sub(r'^[\s]*[-*]\s+', '', line)
            add_bullet_point(doc, text)
            i += 1
            continue
        
        if re.match(r'^[\s]*\d+[\.\)]\s+', line):
            text = re.sub(r'^[\s]*\d+[\.\)]\s+', '', line)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.8)
            add_bold_to_paragraph(p, text)
            i += 1
            continue
        
        add_normal_paragraph(doc, line.strip())
        i += 1
    
    doc.save(docx_path)
    print(f"✅ Saved: {docx_path}")

# Convert both files
base = "/home/ubuntu/echoes/docs/business"
convert_md_to_docx(f"{base}/居家康复延续服务包合作方案.md", f"{base}/居家康复延续服务包合作方案.docx")
convert_md_to_docx(f"{base}/人民网合作路径方案.md", f"{base}/人民网合作路径方案.docx")
