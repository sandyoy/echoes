#!/usr/bin/env python3
"""Convert markdown to properly formatted docx using python-docx.
Handles headings, bold, lists, tables, code blocks, blockquotes."""

import re, sys, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_run_with_format(paragraph, text, bold=False, italic=False, size=None, color=None, font_name="Microsoft YaHei"):
    """Add a run with formatting"""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.font.name = font_name
    # Set East Asian font
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>') 
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    return run

def process_inline_tags(doc, paragraph, text):
    """Process **bold** and emoji markers in text"""
    # Split by **...** pattern
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            add_run_with_format(paragraph, part[2:-2], bold=True)
        else:
            add_run_with_format(paragraph, part)

def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=None):
    """Set paragraph spacing"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = line_spacing

def md_to_docx(input_path, output_path):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # Configure heading styles
    for level, sz, color in [(1, 20, RGBColor(0x1F, 0x4E, 0x79)),
                              (2, 16, RGBColor(0x2E, 0x75, 0xB6)),
                              (3, 14, RGBColor(0x2E, 0x75, 0xB6)),
                              (4, 12, RGBColor(0x33, 0x33, 0x33))]:
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Microsoft YaHei'
        heading_style.font.size = Pt(sz)
        heading_style.font.bold = True
        heading_style.font.color.rgb = color
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    with open(input_path, 'r', encoding='utf-8') as f:
        md_lines = f.read().split('\n')
    
    i = 0
    in_code_block = False
    code_lines = []
    
    while i < len(md_lines):
        line = md_lines[i].strip()
        
        # Code block
        if line.startswith('```'):
            if in_code_block:
                # End code block - add as monospaced paragraph
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                # Add left border for code block
                pPr = p._p.get_or_add_pPr()
                pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="6" w:space="6" w:color="2E75B6"/></w:pBdr>')
                pPr.append(pBdr)
                # Shading
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F5F5F5"/>')
                pPr.append(shading)
                
                for cl in code_lines:
                    if cl.strip():
                        run = p.add_run(cl + '\n')
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(md_lines[i])
            i += 1
            continue
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Horizontal rule
        if line.startswith('---') or line.startswith('***'):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/></w:pBdr>')
            pPr.append(pBdr)
            i += 1
            continue
        
        # Heading 1
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_heading(line[2:], level=1)
            set_paragraph_spacing(p, before=18, after=8)
            i += 1
            continue
        
        # Heading 2
        if line.startswith('## ') and not line.startswith('### '):
            p = doc.add_heading(line[3:], level=2)
            set_paragraph_spacing(p, before=14, after=6)
            i += 1
            continue
        
        # Heading 3
        if line.startswith('### ') and not line.startswith('#### '):
            p = doc.add_heading(line[4:], level=3)
            set_paragraph_spacing(p, before=10, after=4)
            i += 1
            continue
        
        # Heading 4
        if line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
            set_paragraph_spacing(p, before=8, after=4)
            i += 1
            continue
        
        # Blockquote (> text)
        if line.startswith('> '):
            # Collect consecutive blockquote lines
            content = line[2:]
            while i + 1 < len(md_lines) and md_lines[i+1].strip().startswith('> '):
                i += 1
                content += ' ' + md_lines[i].strip()[2:]
            
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=6, after=6)
            pPr = p._p.get_or_add_pPr()
            # Left border
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="12" w:space="8" w:color="2E75B6"/></w:pBdr>')
            pPr.append(pBdr)
            # Light blue background
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F2F7FB"/>')
            pPr.append(shading)
            
            process_inline_tags(doc, p, content)
            i += 1
            continue
        
        # Table (pipe table)
        if line.startswith('|') and '|' in line:
            # Check if next line is separator
            if i + 1 < len(md_lines) and '---' in md_lines[i+1]:
                headers = [h.strip() for h in line.split('|') if h.strip()]
                i += 2  # Skip separator line
                
                # Parse alignments
                sep = md_lines[i-1].strip()
                seps = [s.strip() for s in sep.split('|') if s.strip()]
                aligns = ['left'] * len(headers)
                for j, s in enumerate(seps[:len(headers)]):
                    if s.startswith(':') and s.endswith(':'):
                        aligns[j] = 'center'
                    elif s.endswith(':'):
                        aligns[j] = 'right'
                
                # Read data rows
                rows = []
                while i < len(md_lines) and md_lines[i].strip().startswith('|'):
                    cells = [c.strip() for c in md_lines[i].split('|') if c.strip()]
                    if cells:
                        rows.append(cells[:len(headers)])
                    i += 1
                
                if headers and rows:
                    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    # Style table
                    table.style = 'Table Grid'
                    
                    # Header row
                    for j, h in enumerate(headers):
                        cell = table.rows[0].cells[j]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        run = p.add_run(h)
                        run.bold = True
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.name = 'Microsoft YaHei'
                        set_cell_shading(cell, '2E75B6')
                    
                    # Data rows
                    for ri, row in enumerate(rows):
                        for j, cell_val in enumerate(row):
                            cell = table.rows[ri+1].cells[j]
                            cell.text = ''
                            p = cell.paragraphs[0]
                            process_inline_tags(doc, p, cell_val)
                            for run in p.runs:
                                run.font.size = Pt(10)
                            # Alternate row shading
                            if ri % 2 == 1:
                                set_cell_shading(cell, 'F2F7FB')
                    
                    # Add spacing after table
                    doc.add_paragraph()
                continue
            else:
                # Single row table - just treat as paragraph
                cells = [c.strip() for c in line.split('|') if c.strip()]
                if cells:
                    p = doc.add_paragraph()
                    add_run_with_format(p, ' | '.join(cells))
                i += 1
                continue
        
        # List items (- or * )
        if line.startswith('- ') or line.startswith('* '):
            content = line[2:]
            # Collect continuation lines
            while i + 1 < len(md_lines):
                next_line = md_lines[i + 1].strip()
                if not next_line or next_line.startswith('- ') or next_line.startswith('* ') or next_line.startswith('#') or next_line.startswith('|') or next_line.startswith('> ') or next_line.startswith('```') or next_line.startswith('---'):
                    break
                if next_line.startswith('```') or next_line.startswith('```'):
                    break
                content += ' ' + next_line.strip()
                i += 1
            
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=2, after=2)
            # Add indentation
            pf = p.paragraph_format
            pf.left_indent = Cm(1.0)
            pf.first_line_indent = Cm(-0.5)
            
            # Bullet character
            add_run_with_format(p, '• ', bold=False, size=11, color=RGBColor(0x2E, 0x75, 0xB6))
            process_inline_tags(doc, p, content)
            i += 1
            continue
        
        # Regular paragraph - collect consecutive lines into one paragraph
        para_text = line
        while i + 1 < len(md_lines):
            next_line = md_lines[i + 1].strip()
            if not next_line or next_line.startswith('#') or next_line.startswith('- ') or next_line.startswith('* ') or next_line.startswith('|') or next_line.startswith('> ') or next_line.startswith('```') or next_line.startswith('---') or next_line.startswith('***'):
                break
            para_text += ' ' + next_line.strip()
            i += 1
        
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=2, after=2)
        process_inline_tags(doc, p, para_text)
        i += 1
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    doc.save(output_path)
    return True


if __name__ == '__main__':
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    success = md_to_docx(input_path, output_path)
    if success:
        print(f"Done. Generated {output_path} ({os.path.getsize(output_path)} bytes)")
    else:
        print("Failed")
        sys.exit(1)
