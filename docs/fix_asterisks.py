"""Remove **asterisk** bold markers from docx without breaking paragraph structure"""
from docx import Document
import sys, os

def remove_asterisks_from_docx(input_path, output_path):
    doc = Document(input_path)
    changes = 0
    for p in doc.paragraphs:
        # Process runs: remove ** markers from text
        for run in p.runs:
            if '**' in run.text:
                run.text = run.text.replace('**', '')
                changes += 1
        # Also clean up any doubled-up text (bold markers that split across runs)
        # Merge runs if needed — simple approach: just clean each run
    
    # Also clean tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if '**' in run.text:
                            run.text = run.text.replace('**', '')
                            changes += 1
    
    doc.save(output_path)
    return changes

if __name__ == '__main__':
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    c = remove_asterisks_from_docx(input_path, output_path)
    print(f"Done. Removed ** from {c} runs. Saved to {output_path}")
